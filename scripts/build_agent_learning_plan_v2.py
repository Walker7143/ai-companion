from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "AI_Agent开发系统学习计划_v2_2026.docx"


@dataclass
class WeekPlan:
    week: str
    theme: str
    objective: str
    deliverable: str
    resources: list[str]
    tasks: list[str]
    acceptance: list[str]
    notes: list[str]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name, size, bold, color in [
        ("Title", 24, True, "16324F"),
        ("Heading 1", 16, True, "16324F"),
        ("Heading 2", 13, True, "274C77"),
        ("Heading 3", 11.5, True, "274C77"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = styles["Normal"]
        code_style.font.name = "Consolas"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code_style.font.size = Pt(9.5)


def apply_page_setup(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI Agent 开发系统学习计划 V2")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("面向 Java 后端工程师 · Python + Java 双轨版 · 2026-06 修订版")
    run.bold = True
    run.font.size = Pt(11.5)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "定位：从 0 到能做出可上线的 Agent 服务，不止会跑 Demo，还能理解工程化、评估、可观测性与安全边界。"
    )

    box = doc.add_paragraph()
    box.paragraph_format.space_before = Pt(8)
    box.paragraph_format.space_after = Pt(8)
    for line in [
        "建议周期：12-14 周",
        "建议投入：工作日每天 1-2 小时，周末 3-4 小时",
        "学习原则：先做最小闭环，再加记忆、RAG、MCP、工作流，最后再碰多 Agent",
        "核心产出：1 个最小 Agent、1 个 RAG Agent、1 个 MCP Server、1 个企业级 Java Agent 雏形",
    ]:
        add_bullet(doc, line, level=0, target=box if line == "建议周期：12-14 周" else None)


def add_para(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)


def add_bullet(doc: Document, text: str, level: int = 0, target=None) -> None:
    p = target if target is not None else doc.add_paragraph()
    if target is None:
        p.style = "Normal"
    p.paragraph_format.left_indent = Inches(0.25 + 0.22 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(2)
    p.add_run("• ")
    p.add_run(text)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{idx}. ")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        set_cell_shading(hdr_cells[idx], "D9EAF7")
        for p in hdr_cells[idx].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build_week_plans() -> list[WeekPlan]:
    return [
        WeekPlan(
            week="第 1 周",
            theme="环境准备 + Python 速通 + Java 起跑",
            objective="做到能看懂 Python 教程，能用 Python 和 Java 各打通一次 LLM API。",
            deliverable="`python/translate-api` + `java/translate-api` 两个最小项目。",
            resources=[
                "Python 官方教程：https://docs.python.org/3/tutorial/",
                "FastAPI Tutorial：https://fastapi.tiangolo.com/tutorial/",
                "Spring AI Overview：https://docs.spring.io/spring-ai/reference/",
                "OpenAI Tools Guide：https://developers.openai.com/api/docs/guides/tools",
            ],
            tasks=[
                "安装 Python 3.12+、uv、IDE 插件、Postman/Apifox、JDK 21、Maven/Gradle。",
                "用 Python 官方教程扫完 3、4、5、7、8、9、12 章，重点只看你会马上用到的语法。",
                "写一个 FastAPI 的 GET/POST Demo，体会路由、请求体、Pydantic 校验。",
                "创建 Spring Boot + Spring AI 工程，用 ChatClient 调通一个最简单的问答接口。",
                "把同一条 system prompt、同一条 user prompt 在 Python 和 Java 都跑一遍，观察输出差异。",
            ],
            acceptance=[
                "你能解释 `system/user/assistant` 三类消息分别干什么。",
                "两个翻译接口都能稳定返回结果。",
                "你能看懂 FastAPI 和 Spring Boot 的最小 Controller 代码。",
            ],
            notes=[
                "本周目标不是学会 Python，而是做到“看教程不痛苦”。",
                "不要从 asyncio 深挖开始，先把同步版写通。",
            ],
        ),
        WeekPlan(
            week="第 2 周",
            theme="LLM 基础 + 第一个单 Agent",
            objective="掌握 Token、Prompt、Tool Calling、ReAct 循环，并手写最小 Agent Loop。",
            deliverable="一个原生 Python 版 CLI Agent，带天气、计算器、待办 3 个工具。",
            resources=[
                "OpenAI Practical Guide PDF：https://cdn.openai.com/business-guides-and-resources/a-practical-guide-to-building-agents.pdf",
                "OpenAI Tools Guide：https://developers.openai.com/api/docs/guides/tools",
                "Anthropic Prompt Engineering：https://platform.claude.com/docs/en/build-with-claude/prompt-engineering/overview",
                "Spring AI Tool Calling：https://docs.spring.io/spring-ai/reference/api/tools.html",
            ],
            tasks=[
                "理解 Agent 最小闭环：模型输出工具调用意图 -> 代码执行工具 -> 把结果喂回模型 -> 再决策。",
                "手写 while-loop 版本 Agent，不依赖 LangChain/LangGraph。",
                "给每个工具写清楚描述、参数说明、错误返回格式。",
                "Java 侧用 Spring AI 的 `@Tool` 和 `ChatClient.tools(...)` 跑一个最小示例，不追求复杂场景。",
            ],
            acceptance=[
                "给 Agent 一个复合任务时，它能连续调用 2 个以上工具并整合答案。",
                "你能解释“模型不会自己执行工具，执行工具的是你的应用代码”。",
                "你能说出工具描述写得烂时，模型最容易出哪类错。",
            ],
            notes=[
                "本周是全计划最重要的一周，后面所有框架都只是把这套循环封装起来。",
                "先别碰多 Agent。",
            ],
        ),
        WeekPlan(
            week="第 3 周",
            theme="RAG 最小闭环",
            objective="用 Python 和 Java 各做一版文档问答系统，理解切分、向量化、检索、生成。",
            deliverable="一个可上传 PDF 并提问的 RAG Demo。",
            resources=[
                "Spring AI RAG：https://docs.spring.io/spring-ai/reference/api/retrieval-augmented-generation.html",
                "Spring AI Embeddings：https://docs.spring.io/spring-ai/reference/api/embeddings.html",
                "pgvector：https://github.com/pgvector/pgvector",
                "LangGraph / LangChain 文档入口：https://docs.langchain.com/oss/python/langgraph/overview",
            ],
            tasks=[
                "Python 端用本地 PDF + 向量库做最小问答。",
                "Java 端用 Spring AI + PGVector 做同样流程。",
                "把 chunk 大小改成两到三种配置，观察召回质量变化。",
                "把用户问题和召回片段打印到日志，开始建立可调试意识。",
            ],
            acceptance=[
                "回答里能体现出确实引用了文档内容，而不是模型凭空发挥。",
                "你能解释为什么 chunk 太大或太小都会影响效果。",
                "PGVector 中能看到向量数据入库，且可以执行相似检索。",
            ],
            notes=[
                "别急着做花哨 UI，本周核心是管道打通和效果感知。",
                "Java 侧优先选 PGVector，和现有业务数据库思路更接近。",
            ],
        ),
        WeekPlan(
            week="第 4 周",
            theme="记忆系统 + 对话状态",
            objective="理解短期记忆、长期记忆、事实记忆的职责差异，并做出可跨进程恢复的对话 Agent。",
            deliverable="支持重启后记住用户偏好的对话 Agent。",
            resources=[
                "Spring AI Chat Memory：https://docs.spring.io/spring-ai/reference/api/chat-memory.html",
                "OpenAI Evals Guide：https://developers.openai.com/api/docs/guides/evals",
                "OpenAI Practical Guide PDF：https://cdn.openai.com/business-guides-and-resources/a-practical-guide-to-building-agents.pdf",
            ],
            tasks=[
                "把对话历史、摘要、用户事实拆成三层，避免把所有东西都塞进上下文。",
                "为“记住用户偏好”设计单独的数据结构，而不是每轮全文喂模型。",
                "实现一个简单的摘要压缩策略，避免上下文无限增长。",
                "写 10 条记忆回归测试问题，验证恢复效果。",
            ],
            acceptance=[
                "重启服务后，Agent 还能说出用户姓名、偏好、最近一个待办。",
                "你能区分“会话历史保留”和“长期事实提取”这两件事。",
            ],
            notes=[
                "很多项目所谓记忆，其实只是把历史消息原样拼接，这不算真正的长期记忆。",
            ],
        ),
        WeekPlan(
            week="第 5 周",
            theme="Workflow First：先学可控工作流，再学大框架",
            objective="建立比“直接上多 Agent”更稳的工程心智：路由、并行、审核、编排者-执行者模式。",
            deliverable="一个包含分类、路由、并行步骤的小工作流。",
            resources=[
                "Spring AI Building Effective Agents：https://docs.spring.io/spring-ai/reference/api/effective-agents.html",
                "OpenAI Practical Guide PDF：https://cdn.openai.com/business-guides-and-resources/a-practical-guide-to-building-agents.pdf",
            ],
            tasks=[
                "分别实现 4 种小模式：分类路由、并行调用、审核后重试、orchestrator-workers。",
                "每种模式先用流程图画出来，再写代码。",
                "把“Agent 像工作流系统”这个认知落地，而不是停留在概念层。",
            ],
            acceptance=[
                "你能解释什么场景该用单 Agent，什么场景适合 workflow，什么场景才需要多 Agent。",
                "你能把一个客服问题拆成分类 -> 查询 -> 决策 -> 输出这四步。",
            ],
            notes=[
                "这一周是 V2 新增的关键阶段，用来避免直接跳到 LangGraph 后只会拼 API。",
            ],
        ),
        WeekPlan(
            week="第 6 周",
            theme="LangGraph 深入 + Java 侧 Tool/Advisor 心智",
            objective="掌握 LangGraph 的 State、Node、Edge、Checkpoint，同时理解 Java 侧为什么更偏高层封装。",
            deliverable="一个客服或报销审批 Agent 工作流。",
            resources=[
                "LangGraph Overview：https://docs.langchain.com/oss/python/langgraph/overview",
                "Spring AI Tool Calling：https://docs.spring.io/spring-ai/reference/api/tools.html",
                "Spring AI Chat Client API：https://docs.spring.io/spring-ai/reference/api/chatclient.html",
            ],
            tasks=[
                "用 LangGraph 实现状态对象在多个节点间传递。",
                "加入人工确认节点，如退款超额需要审批。",
                "Java 侧用 ChatClient + Tool Calling + Memory 实现同场景简化版。",
            ],
            acceptance=[
                "中断后能恢复流程状态。",
                "你能解释 StateGraph 和传统状态机/工作流引擎的相似点。",
            ],
            notes=[
                "Java 侧本周不用硬追图编排对等能力，核心是建立映射关系。",
            ],
        ),
        WeekPlan(
            week="第 7 周",
            theme="工具设计专项",
            objective="把工具设计做成你的强项，学会把 Java 业务服务拆成让模型易用的工具接口。",
            deliverable="6-8 个高质量工具定义，含参数说明、错误语义、权限边界。",
            resources=[
                "Spring AI Tool Calling：https://docs.spring.io/spring-ai/reference/api/tools.html",
                "MCP Architecture：https://modelcontextprotocol.io/docs/learn/architecture",
                "OpenAI Tools Guide：https://developers.openai.com/api/docs/guides/tools",
            ],
            tasks=[
                "把订单、用户、退款、消息通知等业务接口包装成清晰工具。",
                "为每个工具写“何时该用 / 何时不该用”。",
                "统一错误格式，例如 `USER_NOT_FOUND`、`ORDER_NOT_REFUNDABLE`。",
                "给高风险工具补人工审批或二次确认机制。",
            ],
            acceptance=[
                "模型在 10 次调用里，大多数时候能选对工具。",
                "错误信息能帮助模型修正参数，而不是只打印异常栈。",
            ],
            notes=[
                "这一周的收获会直接决定你后面做业务 Agent 的质量上限。",
            ],
        ),
        WeekPlan(
            week="第 8 周",
            theme="MCP 协议入门到实战",
            objective="搞懂 MCP 的 Host / Client / Server、tools/list、tools/call、resources、prompts。",
            deliverable="一个 Python MCP Server + 一个 Java MCP Server。",
            resources=[
                "MCP Architecture：https://modelcontextprotocol.io/docs/learn/architecture",
                "Spring AI MCP Getting Started：https://docs.spring.io/spring-ai/reference/guides/getting-started-mcp.html",
                "Spring AI MCP Overview：https://docs.spring.io/spring-ai/reference/api/mcp/mcp-overview.html",
            ],
            tasks=[
                "Python 端把已有工具封成 MCP Server。",
                "Java 端用 `@McpTool` / `@McpToolParam` 做 MCP Server，而不是继续沿用旧写法。",
                "验证同一个 Client 能连接不同语言写的 MCP Server。",
                "理解 MCP 不是“又一个 Agent 框架”，而是工具和上下文协议。",
            ],
            acceptance=[
                "你能解释 `tools/list` 与 `tools/call` 的往返过程。",
                "两个 MCP Server 都能被客户端发现并调用。",
            ],
            notes=[
                "V2 明确改成当前 Spring AI 文档推荐的注解和 starter。",
            ],
        ),
        WeekPlan(
            week="第 9 周",
            theme="RAG 效果工程",
            objective="从“能检索”进化到“答得稳”，掌握 chunk、metadata、hybrid、rerank、引用返回。",
            deliverable="一个效果明显优于第 3 周版本的 RAG Agent。",
            resources=[
                "OpenAI Evals Guide：https://developers.openai.com/api/docs/guides/evals",
                "pgvector：https://github.com/pgvector/pgvector",
                "Langfuse Docs：https://langfuse.com/docs",
            ],
            tasks=[
                "比较不同 chunk 策略和 overlap。",
                "给文档加 metadata，例如分类、日期、业务线。",
                "补充引用返回，让答案可追溯到片段或文档位置。",
                "建立 20 条离线问答集做回归。",
            ],
            acceptance=[
                "你能说明某个效果提升是来自哪一项改动，而不是凭感觉。",
                "RAG 回答能带上来源片段或文档标识。",
            ],
            notes=[
                "这一周开始，重点从“写功能”切到“调系统”。",
            ],
        ),
        WeekPlan(
            week="第 10 周",
            theme="评估体系 + Prompt 管理",
            objective="建立最小 Evals 体系，让 Agent 的质量可以回归，而不是每次都手测。",
            deliverable="一个带黄金测试集、LLM-as-Judge、Prompt 版本管理的仓库。",
            resources=[
                "OpenAI Evals Guide：https://developers.openai.com/api/docs/guides/evals",
                "Spring AI LLM-as-a-Judge Guide：https://docs.spring.io/spring-ai/reference/",
                "Anthropic Prompt Engineering：https://platform.claude.com/docs/en/build-with-claude/prompt-engineering/overview",
            ],
            tasks=[
                "准备 20-50 条真实业务问题，手工定义通过标准。",
                "Prompt、Tool Description、System Policy 全部入 Git。",
                "增加 LLM-as-Judge，辅助判断回答质量。",
                "把“输出逐字一致”改成“行为和约束一致”。",
            ],
            acceptance=[
                "任何改动后都能跑一轮最小回归，不靠拍脑袋。",
                "你能说出哪类问题适合规则判断，哪类适合 LLM-as-Judge。",
            ],
            notes=[
                "这周会让你真正有“工程感”。",
            ],
        ),
        WeekPlan(
            week="第 11 周",
            theme="Observability + Guardrails + 安全",
            objective="学会观测 Token、延迟、工具调用、错误率，处理 Prompt Injection 与高风险动作审批。",
            deliverable="一个带 trace、metrics、审批点的 Java Agent 服务。",
            resources=[
                "Spring AI Observability：https://docs.spring.io/spring-ai/reference/observability/",
                "Spring AI MCP Security：https://docs.spring.io/spring-ai/reference/api/mcp/mcp-security.html",
                "OpenAI Safety Best Practices：https://developers.openai.com/api/docs/guides/safety-best-practices",
            ],
            tasks=[
                "接入 Micrometer/Prometheus，观察 ChatClient 和向量检索指标。",
                "记录 tool call 入参、耗时、失败原因，但注意脱敏。",
                "设计高风险动作审批，如退款、发券、删数据。",
                "处理最基本的 prompt injection 风险：输入清洗、工具最小暴露、只读/写分离。",
            ],
            acceptance=[
                "你能在图表上看到调用量、耗时、错误率。",
                "高风险动作不会仅凭模型一句话直接落库执行。",
            ],
            notes=[
                "很多 Demo 死在这里：没有 trace，就调不动；没有 guardrails，就不敢上线。",
            ],
        ),
        WeekPlan(
            week="第 12 周",
            theme="多 Agent：只在必要时使用",
            objective="理解 manager-worker 模式，能判断多 Agent 什么时候真有必要。",
            deliverable="一个多 Agent 报告生成或客服分工系统。",
            resources=[
                "OpenAI Practical Guide PDF：https://cdn.openai.com/business-guides-and-resources/a-practical-guide-to-building-agents.pdf",
                "Spring AI Building Effective Agents：https://docs.spring.io/spring-ai/reference/api/effective-agents.html",
                "LangGraph Overview：https://docs.langchain.com/oss/python/langgraph/overview",
            ],
            tasks=[
                "先实现 manager-worker，再考虑更复杂的 handoff 方案。",
                "让 manager 控制用户入口和最终整合，worker 只做专业任务。",
                "比较单 Agent、workflow、多 Agent 的复杂度差异。",
            ],
            acceptance=[
                "你能说出这个场景为什么需要多 Agent，而不是单 Agent + 更多工具。",
                "每个 worker 职责边界明确，不相互串话。",
            ],
            notes=[
                "V2 保留多 Agent，但明确把它放到后面，而不是中前期重点。",
            ],
        ),
        WeekPlan(
            week="第 13-14 周",
            theme="企业级落地项目",
            objective="把前面所有能力收敛成一个 Java 主导的实战项目。",
            deliverable="一个可写进简历、可演示、可扩展的 Agent 项目。",
            resources=[
                "Spring AI Samples：https://github.com/spring-projects/spring-ai/tree/main/spring-ai-samples",
                "Langfuse Docs：https://langfuse.com/docs",
                "pgvector：https://github.com/pgvector/pgvector",
            ],
            tasks=[
                "推荐项目 1：客服退款审批 Agent。",
                "推荐项目 2：内部知识库问答 + 工单创建 Agent。",
                "推荐项目 3：销售/运营 Copilot，支持查询、总结、通知、人工审批。",
                "输出架构图、接口说明、评估结果、trace 截图、演示脚本。",
            ],
            acceptance=[
                "项目具备真实业务边界：鉴权、日志、回归、审批、异常处理。",
                "你可以向面试官清楚讲出：为什么选 Java 做主框架，Python 在哪里辅助。",
            ],
            notes=[
                "如果时间有限，宁可把 1 个项目做深，也不要做 3 个半成品。",
            ],
        ),
    ]


def add_week_section(doc: Document, plan: WeekPlan) -> None:
    doc.add_heading(f"{plan.week}：{plan.theme}", level=2)
    add_para(doc, f"学习目标：{plan.objective}", bold_prefix="学习目标：")
    add_para(doc, f"本周产出：{plan.deliverable}", bold_prefix="本周产出：")

    doc.add_heading("建议任务", level=3)
    add_numbered(doc, plan.tasks)

    doc.add_heading("本周验收", level=3)
    for item in plan.acceptance:
        add_bullet(doc, item)

    doc.add_heading("必看资料", level=3)
    for item in plan.resources:
        add_bullet(doc, item)

    doc.add_heading("提醒", level=3)
    for item in plan.notes:
        add_bullet(doc, item)


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    apply_page_setup(doc)
    add_title(doc)

    doc.add_heading("一、这版 V2 相比旧版新增了什么", level=1)
    for item in [
        "把 Spring AI 的旧 API 表述改成当前稳定文档路径，尤其是 `EmbeddingModel`、`Tool Calling`、`MCP`、`Observability`。",
        "新增 `Workflow First` 阶段，避免一开始就跳进多 Agent。",
        "新增 `RAG 效果工程`、`Evals`、`Observability`、`Guardrails` 四块，让学习结果更接近真实上线能力。",
        "把每周计划扩成可执行清单：目标、任务、验收、资料、提醒。",
        "统一强调一条主线：Python 用来快学和快验证，Java 用来做最终工程产出。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("二、你最适合的学习策略", level=1)
    add_para(
        doc,
        "你是 Java 后端工程师，最容易犯的两个错误是：第一，想把所有东西都用 Java 一步到位；第二，一上来就学多 Agent、复杂编排。V2 的原则是：用 Python 快速建立概念和直觉，用 Java 承接最终工程化成果。",
    )
    add_table(
        doc,
        ["场景", "建议语言", "原因"],
        [
            ["学新概念", "Python 优先", "教程多、示例新、反馈快"],
            ["做正式服务", "Java 优先", "你已有鉴权、事务、监控、部署经验"],
            ["做工具/MCP Server", "Python 和 Java 都做一遍", "一遍看协议，一遍看工程化"],
            ["做复杂编排", "先 Python 理解，再决定是否 Java 落地", "LangGraph 思想最清楚"],
            ["做简历项目", "Java 主体 + Python 辅助", "更贴近你的职业优势"],
        ],
    )

    doc.add_heading("三、学习总路线图", level=1)
    add_table(
        doc,
        ["阶段", "时间", "主题", "关键产出"],
        [
            ["0", "第 1 周", "环境准备 + Python 速通", "Python/Java 各一个最小 LLM API"],
            ["1", "第 2 周", "LLM 基础 + 单 Agent", "CLI Agent 最小闭环"],
            ["2", "第 3-4 周", "RAG + 记忆", "文档问答 + 长期记忆 Agent"],
            ["3", "第 5-7 周", "Workflow + LangGraph + 工具设计", "审批/客服工作流 Agent"],
            ["4", "第 8 周", "MCP", "Python MCP Server + Java MCP Server"],
            ["5", "第 9-11 周", "RAG 效果工程 + Evals + Guardrails", "可回归、可观测的 Agent 服务"],
            ["6", "第 12-14 周", "多 Agent + 企业级落地", "可演示、可写简历的完整项目"],
        ],
    )

    doc.add_heading("四、先建立正确心智模型", level=1)
    for item in [
        "Agent 不是更聪明的聊天机器人，而是“会推理的工作流系统”。",
        "Tool Calling 不是模型替你执行代码，而是模型提出调用请求，你的应用负责真正执行。",
        "RAG 的难点不是“向量库接上没有”，而是效果工程：切分、召回、过滤、重排、引用、评估。",
        "多 Agent 不是默认高级方案。单 Agent + 好工具 + 清晰工作流，通常先够用。",
        "企业落地真正拉开差距的是：权限、日志、评估、可观测性、人工审批、安全边界。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("五、每周详细学习路线", level=1)
    for plan in build_week_plans():
        add_week_section(doc, plan)

    page_break(doc)
    doc.add_heading("六、三个推荐实战项目", level=1)
    projects = [
        (
            "项目 A：客服退款审批 Agent",
            [
                "适合度最高，因为它天然包含查询、规则判断、工具调用、人工审批、日志追踪。",
                "Java 侧负责订单、退款、通知、权限；Python 或 Java Agent 侧负责决策编排。",
                "最适合展示你把传统后端能力迁移到 Agent 场景的价值。",
            ],
        ),
        (
            "项目 B：企业知识库问答 + 工单创建 Agent",
            [
                "包含 RAG、知识检索、来源引用、问题分类、工单落地。",
                "适合展示 RAG 效果调优、工具设计、评估体系。",
            ],
        ),
        (
            "项目 C：运营/销售 Copilot",
            [
                "让 Agent 查询用户、订单、活动、日报，并能生成摘要、提醒和操作建议。",
                "适合展示 MCP、工具聚合、观测和权限边界。",
            ],
        ),
    ]
    for title, bullets in projects:
        doc.add_heading(title, level=2)
        for item in bullets:
            add_bullet(doc, item)

    doc.add_heading("七、如果每天只有 1 小时，怎么学不枯燥", level=1)
    add_table(
        doc,
        ["时间段", "做什么", "目的"],
        [
            ["15 分钟", "看 1 篇官方文档/教程", "补正确认知"],
            ["30 分钟", "写一点代码或跑一个实验", "形成肌肉记忆"],
            ["15 分钟", "记录今天的坑和结论", "沉淀成自己的方法论"],
        ],
    )
    for item in [
        "每周至少保留 1 个“可见产出”，比如一个接口、一个流程图、一组 eval 数据。",
        "不要连续 3 天只看文档不写代码，这会极大增加枯燥感。",
        "不要把“刷 20 篇公众号”误当成学习进度，Agent 学习一定要靠做闭环。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("八、你应该重点收藏的网站和文章", level=1)
    add_table(
        doc,
        ["类别", "建议站点 / 文章", "为什么值得看"],
        [
            ["OpenAI Agent 基础", "https://cdn.openai.com/business-guides-and-resources/a-practical-guide-to-building-agents.pdf", "讲清楚单 Agent、多 Agent、工具、guardrails 的基础框架"],
            ["OpenAI 工具调用", "https://developers.openai.com/api/docs/guides/tools", "理解 Tool Calling 的官方定义和实践"],
            ["OpenAI 评估", "https://developers.openai.com/api/docs/guides/evals", "补上 Evals 心智，避免只凭主观判断质量"],
            ["Spring AI 总文档", "https://docs.spring.io/spring-ai/reference/", "Java 主线资料入口"],
            ["Spring AI Tool Calling", "https://docs.spring.io/spring-ai/reference/api/tools.html", "Java 工具设计和调用的核心页面"],
            ["Spring AI Effective Agents", "https://docs.spring.io/spring-ai/reference/api/effective-agents.html", "补 workflow 和多模式编排心智"],
            ["Spring AI MCP", "https://docs.spring.io/spring-ai/reference/guides/getting-started-mcp.html", "当前 Java MCP 入门最佳入口"],
            ["Spring AI Chat Memory", "https://docs.spring.io/spring-ai/reference/api/chat-memory.html", "Java 对话记忆主线"],
            ["Spring AI Observability", "https://docs.spring.io/spring-ai/reference/observability/", "生产化必备"],
            ["Python 官方教程", "https://docs.python.org/3/tutorial/", "最快速的程序员友好型 Python 入门"],
            ["FastAPI Tutorial", "https://fastapi.tiangolo.com/tutorial/", "Python Web 原型首选"],
            ["LangGraph Overview", "https://docs.langchain.com/oss/python/langgraph/overview", "工作流式 Agent 编排的核心概念入口"],
            ["MCP 架构", "https://modelcontextprotocol.io/docs/learn/architecture", "把 MCP 当协议来理解，而不是当框架来理解"],
            ["pgvector", "https://github.com/pgvector/pgvector", "Java 后端做 RAG 时最实用的向量存储基础"],
            ["Langfuse", "https://langfuse.com/docs", "观测、trace、提示词与评估管理"],
            ["Anthropic Prompt Engineering", "https://platform.claude.com/docs/en/build-with-claude/prompt-engineering/overview", "prompt 与评估关系讲得比较稳"],
        ],
    )

    doc.add_heading("九、V2 特别强调的避坑指南", level=1)
    for item in [
        "不要把“结果文字一模一样”当作跨语言验收标准，应该看行为是否正确、工具是否选对、约束是否满足。",
        "不要把全部历史消息都塞给模型冒充记忆系统。",
        "不要在没有 trace 的前提下做复杂调优，你会不知道问题出在 prompt、工具、检索还是数据。",
        "不要把 MCP 误解成 LangChain 的替代品，它是协议层，不是完整编排层。",
        "不要一开始就追多 Agent，先把单 Agent + 工具 + 工作流打透。",
        "不要把高风险写操作直接暴露给模型，必须有权限校验和人工审批。",
        "不要把公众号二手文章当作长期主资料，尤其是 Spring AI 和 OpenAI 的 API 名称变化很快。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("十、最终学习验收标准", level=1)
    add_numbered(
        doc,
        [
            "你能手写一个最小 Agent Loop，并解释每一步为什么存在。",
            "你能用 Java 做出带 Tool Calling、Memory、RAG 的 Agent 服务。",
            "你能把已有业务接口包装成模型更容易调用的工具。",
            "你能搭一个 MCP Server，并解释它和普通 REST API 的区别。",
            "你能建立最小 Evals 与可观测性体系，而不是只做 Demo。",
            "你能判断什么时候该用单 Agent、Workflow、MCP、多 Agent。",
            "你能拿出一个企业级实战项目，讲清楚架构、约束、收益和风险。",
        ],
    )

    doc.add_heading("附录 A：推荐的仓库组织方式", level=1)
    add_para(doc, "如果你要边学边做项目，建议直接用下面的目录结构。")
    code = [
        "agent-learning/",
        "  python/",
        "    week02-single-agent/",
        "    week03-rag/",
        "    week08-mcp-server/",
        "  java/",
        "    week01-translate-api/",
        "    week03-rag-spring-ai/",
        "    week08-mcp-server/",
        "    week13-enterprise-agent/",
        "  prompts/",
        "  evals/",
        "  docs/",
        "    architecture/",
        "    traces/",
    ]
    for line in code:
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(line)

    doc.add_heading("附录 B：Java 侧当前建议优先掌握的 Spring AI 能力", level=1)
    add_table(
        doc,
        ["能力", "当前主 API / 文档入口", "学习优先级"],
        [
            ["基础问答", "ChatClient", "最高"],
            ["工具调用", "@Tool / tools() / ToolCallback", "最高"],
            ["Embedding", "EmbeddingModel", "高"],
            ["记忆", "ChatMemory", "高"],
            ["RAG", "Retrieval Augmented Generation", "高"],
            ["MCP", "@McpTool / MCP starters", "高"],
            ["可观测性", "Observability + Micrometer", "高"],
            ["复杂编排", "Effective Agents / 外部工作流框架", "中"],
        ],
    )

    doc.add_heading("附录 C：建议你每周输出一页复盘", level=1)
    for item in [
        "这周学会了什么？",
        "这周最关键的 1 个误区是什么？",
        "哪一个官方页面最值得重看？",
        "哪个实验最有收获？",
        "如果下周只能保留 1 个目标，它是什么？",
    ]:
        add_bullet(doc, item)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
